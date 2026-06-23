"""Multi-format file loaders: Excel, CSV, PDF, DOCX, TXT, Markdown."""
from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Any

import pandas as pd

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".xlsx", ".xlsm", ".xls", ".csv", ".tsv", ".pdf", ".docx", ".txt", ".md"}


def load_file(file_path: str | Path) -> dict[str, Any]:
    """Detect type and load a single file into a uniform dict.

    Returns:
        {
            "path": str,
            "type": "tabular" | "document" | "text",
            "sheets": {sheet_name: pd.DataFrame} | None,
            "text": str | None,
            "metadata": {rows, columns, sheet_names, ...}
        }
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(str(path))

    ext = path.suffix.lower()
    logger.info("Loading %s (%s) …", path.name, ext)

    if ext in (".xlsx", ".xlsm", ".xls"):
        return _load_excel(path)
    elif ext in (".csv", ".tsv"):
        return _load_csv(path)
    elif ext == ".pdf":
        return _load_pdf(path)
    elif ext == ".docx":
        return _load_docx(path)
    elif ext in (".txt", ".md"):
        return _load_text(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def load_files(paths: list[str]) -> list[dict[str, Any]]:
    return [load_file(p) for p in paths]


# ── Private loaders ──────────────────────────────────────────────────────────

def _load_excel(path: Path) -> dict[str, Any]:
    xl = pd.ExcelFile(path, engine="openpyxl")
    sheets = {}
    total_rows = 0
    for name in xl.sheet_names:
        df = xl.parse(name)
        sheets[name] = df
        total_rows += len(df)
    return {
        "path": str(path),
        "type": "tabular",
        "sheets": sheets,
        "text": None,
        "metadata": {
            "sheet_names": xl.sheet_names,
            "total_rows": total_rows,
            "total_columns": max((len(df.columns) for df in sheets.values()), default=0),
        },
    }


def _load_csv(path: Path) -> dict[str, Any]:
    sep = "\t" if path.suffix == ".tsv" else ","
    df = pd.read_csv(path, sep=sep)
    return {
        "path": str(path),
        "type": "tabular",
        "sheets": {"data": df},
        "text": None,
        "metadata": {"total_rows": len(df), "total_columns": len(df.columns)},
    }


def _load_pdf(path: Path) -> dict[str, Any]:
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages_text = []
            tables = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
                for t in page.extract_tables():
                    if t:
                        tables.append(t)
            full_text = "\n\n".join(pages_text)
    except Exception:
        # Fallback to pypdf
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages_text = [page.extract_text() or "" for page in reader.pages]
        full_text = "\n\n".join(pages_text)
        tables = []

    return {
        "path": str(path),
        "type": "document",
        "sheets": None,
        "text": full_text,
        "metadata": {"pages": len(pages_text) if pages_text else 0, "tables_found": len(tables)},
    }


def _load_docx(path: Path) -> dict[str, Any]:
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for t in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in t.rows]
        tables.append(rows)
    return {
        "path": str(path),
        "type": "document",
        "sheets": None,
        "text": "\n\n".join(paragraphs),
        "metadata": {"paragraphs": len(paragraphs), "tables": len(tables)},
    }


def _load_text(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8")
    return {
        "path": str(path),
        "type": "text",
        "sheets": None,
        "text": text,
        "metadata": {"chars": len(text), "lines": text.count("\n") + 1},
    }
